#!/usr/bin/env python3
"""Convert resume markdown to DOCX with proper formatting."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re, sys

def main():
    md_path = sys.argv[1]
    out_path = sys.argv[2]
    
    with open(md_path) as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    
    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    for line in lines:
        line = line.rstrip('\n')
        
        if line.startswith('# ') and not line.startswith('## '):
            # Name / title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(16)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].upper())
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            # Add bottom border
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(10.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith('- '):
            text = line[2:]
            # Handle bold within bullets
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            p.text = text
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)
        elif line.startswith('---'):
            continue
        elif line.strip() == '':
            continue
        else:
            # Contact info or plain text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = doc.add_paragraph(text)
            if '[placeholder' in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
    
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == '__main__':
    main()
