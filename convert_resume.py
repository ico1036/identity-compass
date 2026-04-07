from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Jiwoong Kim')
run.bold = True
run.font.size = Pt(18)

# Contact
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[PLACEHOLDER Email] | [PLACEHOLDER Phone] | [PLACEHOLDER LinkedIn] | Seoul, South Korea')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(100, 100, 100)

def add_section_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)
    # Add bottom border via paragraph format
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '003366'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)

# SUMMARY
add_section_heading('Summary')
add_body('Quantitative researcher with a PhD in Physics and hands-on experience deploying ML models to live financial markets. PhD research focused on CNN-based pattern recognition for signal extraction from noisy, high-dimensional particle physics data on a 1024-node HPC cluster. Currently developing and operating ML-driven systematic equity strategies (34B KRW AUM) and building agentic AI pipelines that autonomously generate, backtest, and rank 500+ strategy candidates daily. Seeking to bring quantitative rigor, statistical modeling, and applied ML expertise to Citadel Securities.')

# EXPERIENCE
add_section_heading('Professional Experience')
p = add_body('Alpha Bridge ', bold_prefix=None)
p.runs[0].bold = True
p.runs[0].font.size = Pt(11)
run = p.add_run('(formerly Asset Plus Asset Management)')
run.font.size = Pt(10)
run.italic = True

p = doc.add_paragraph()
run = p.add_run('Quantitative Developer / Researcher')
run.bold = True
run.font.size = Pt(10.5)
run = p.add_run(' | Seoul, South Korea | Jan 2024 – Present')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

add_bullet(' Developed and deployed ML models for the S&P500 Growth Fund (34B KRW AUM, +43.54% cumulative return). Automated signal generation, feature engineering from time-series and cross-sectional data, and portfolio rebalancing.', 'ML Signal Generation (Live):')
add_bullet(' Built LLM-based pipelines to extract structured signals from financial text — earnings transcripts, research reports, and news flow. Integrated NLP-derived features into quantitative factor models.', 'NLP / LLM Research Pipeline:')
add_bullet(' Architected an autonomous research system that generates hypotheses, constructs backtests, and evaluates 500+ strategy candidates per day using statistical validation (walk-forward, multiple-testing correction). Led end-to-end pipeline design and infrastructure.', 'Agentic Strategy Discovery:')
add_bullet('Applied rigorous statistical testing (Sharpe ratio significance, drawdown analysis, regime detection) to avoid overfitting and ensure strategy robustness.')

# EDUCATION
add_section_heading('Education')
p = add_body('')
run = p.add_run('Ph.D. in Physics')
run.bold = True
run.font.size = Pt(11)
run = p.add_run(' — Yonsei University, Seoul, South Korea | 2023')
run.font.size = Pt(10)

add_bullet('CNN-based pattern recognition for signal/background classification in LHC experimental data.')
add_bullet('Achieved 1.85× improvement in signal detection efficiency over conventional methods.')
add_bullet('Designed data pipelines processing terabyte-scale datasets on a 1024-node distributed HPC cluster.')
add_bullet('Applied Bayesian inference and hypothesis testing to quantify statistical significance.')
add_bullet('Publication: J. Phys.: Conf. Ser. 2438 (2023).')

p = add_body('')
run = p.add_run('B.S. in Physics')
run.bold = True
run.font.size = Pt(11)
run = p.add_run(' — [PLACEHOLDER University]')
run.font.size = Pt(10)

# SKILLS
add_section_heading('Technical Skills')
skills = [
    ('Languages: ', 'Python (primary), SQL, Bash'),
    ('ML / Deep Learning: ', 'PyTorch, TensorFlow, scikit-learn, CNNs, transformer architectures'),
    ('NLP / LLM: ', 'LLM fine-tuning, prompt engineering, RAG pipelines, agentic AI frameworks'),
    ('Statistics: ', 'Bayesian inference, hypothesis testing, time-series analysis, Monte Carlo methods'),
    ('Infrastructure: ', 'Docker, HPC (MPI/SLURM), distributed computing, PostgreSQL'),
    ('Quantitative Finance: ', 'Factor models, backtesting, walk-forward validation, risk metrics'),
]
for bold, text in skills:
    add_bullet(text, bold)

# AWARDS
add_section_heading('Awards & Publications')
add_bullet('Best Presentation Award, Korean Society for Computational Science and Engineering (KSCSE), 2021')
add_bullet('Outstanding Poster Award, Korean Physical Society (KPS), 2020')
add_bullet('Kim, J. et al., "CNN-based signal classification in high-energy physics," J. Phys.: Conf. Ser. 2438 (2023)')

# ADDITIONAL
add_section_heading('Additional Information')
add_bullet('Open to relocation to Singapore.')
add_bullet('Fluent in Korean; professional proficiency in English.')
add_bullet('[PLACEHOLDER: visa/work authorization status for Singapore]')

doc.save('resume_citadel.docx')
print('Done: resume_citadel.docx')
